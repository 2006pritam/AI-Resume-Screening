import io
from typing import Optional

def extract_text_from_pdf(file_bytes: bytes) -> str:
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(file_bytes))
        text_parts = []
        for page in reader.pages:
            t = page.extract_text()
            if t:
                text_parts.append(t)
        return "\n".join(text_parts).strip()
    except Exception as e:
        try:
            return file_bytes.decode('utf-8', errors='ignore')
        except Exception:
            return ""

def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        import docx
        doc = docx.Document(io.BytesIO(file_bytes))
        text_parts = []
        for p in doc.paragraphs:
            if p.text.strip():
                text_parts.append(p.text.strip())
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text.strip())
        return "\n".join(text_parts).strip()
    except Exception as e:
        try:
            return file_bytes.decode('utf-8', errors='ignore')
        except Exception:
            return ""

def extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    fn_lower = filename.lower()
    if fn_lower.endswith('.pdf'):
        return extract_text_from_pdf(file_bytes)
    elif fn_lower.endswith(('.docx', '.doc')):
        return extract_text_from_docx(file_bytes)
    else:
        try:
            return file_bytes.decode('utf-8')
        except UnicodeDecodeError:
            return file_bytes.decode('latin-1', errors='ignore')
